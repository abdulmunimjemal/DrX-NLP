import io
import os
from typing import Union
from docx import Document
import fitz  
import pandas as pd
import camelot
from src.core.core import logger

def process_file(file_path: str) -> list:
    """Process different file formats and extract text with tables"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # extract file name and add it as source
    file_name = os.path.basename(file_path)
    
    try:
        if file_path.endswith('.docx'):
            results = process_docx(file_path)
        elif file_path.endswith('.pdf'):
            results = process_pdf(file_path)
        elif file_path.endswith(('.csv', '.xlsx', '.xls', '.xlsm')):
            results = process_spreadsheet(file_path)
        
        for item in results:
                item['source'] = file_name
        return results
    except Exception as e:
        logger.error(f"Error processing {file_path}: {str(e)}")
        return []

def process_docx(file_path: str) -> list:
    doc = Document(file_path)
    content = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # For uniformity, return as a dictionary
            content.append({"type": "paragraph", "text": para.text})
    
    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            # Each cell is processed to grab its text; this follows the current python-docx API.
            row_data = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(row_data))
        if table_rows:
            content.append({"type": "table", "text": "\n".join(table_rows)})
    
    return content

def process_pdf(file_path: str) -> list:
    content = []
    # Open the PDF file using PyMuPDF (fitz)
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc):
            # Extract text explicitly as text
            text = page.get_text("text")
            if text.strip():
                content.append({"page": page_num + 1, "type": "paragraph", "text": text})
            
            # Extract tables with Camelot; you can optionally specify flavor='stream' or 'lattice'
            tables = camelot.read_pdf(file_path, pages=str(page_num + 1))
            for table in tables:
                table_csv = table.df.to_csv(index=False)
                content.append({"page": page_num + 1, "type": "table", "text": table_csv})
    
    return content

def process_spreadsheet(file_path: str) -> list:
    if file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    else:
        # For Excel files, you can let pandas auto-detect the engine or specify 'openpyxl'
        df = pd.read_excel(file_path, engine='openpyxl')
    
    output = io.StringIO()
    df.to_csv(output, index=False)
    # Return as a dictionary for consistency
    return [{"type": "spreadsheet", "text": output.getvalue()}]